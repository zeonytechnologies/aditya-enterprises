from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Styles
title_style = document.styles['Title']
title_font = title_style.font
title_font.name = 'Calibri'
title_font.size = Pt(26)
title_font.bold = True
title_font.color.rgb = RGBColor(0, 51, 153)

heading_style = document.styles['Heading 1']
heading_font = heading_style.font
heading_font.name = 'Calibri'
heading_font.size = Pt(16)
heading_font.color.rgb = RGBColor(0, 102, 204)

heading2_style = document.styles['Heading 2']
heading2_font = heading2_style.font
heading2_font.name = 'Calibri'
heading2_font.size = Pt(13)
heading2_font.color.rgb = RGBColor(0, 0, 0)
heading2_font.bold = True

normal_style = document.styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Calibri'
normal_font.size = Pt(11)

# Title
title = document.add_paragraph('Project Handover Document', style='Title')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = document.add_paragraph('Aditya Enterprises E-Commerce Platform')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True
subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
document.add_paragraph()

document.add_heading('1. Overview', level=1)
document.add_paragraph('This document outlines the complete workflow of the Aditya Enterprises B2B and B2C E-Commerce platform. It is designed to serve as a comprehensive guide for the system administrator and the end customer to understand the core functionalities and day-to-day operations of the platform.')

document.add_heading('2. Customer Workflow (Front-End)', level=1)

document.add_heading('2.1 Account Creation & Login', level=2)
document.add_paragraph('Customers can securely log into the platform using email/password authentication or username-based login. The secure session ensures all browsing and transactions remain private to the logged-in user.', style='List Bullet')

document.add_heading('2.2 Browsing & Navigation', level=2)
document.add_paragraph('Homepage: Customers are greeted with a sliding top banner showcasing active promotions, a dynamic hero section for featured schemes, and a responsive navigation bar.', style='List Bullet')
document.add_paragraph('Catalog: Customers can browse through dedicated sections for Brands, Categories, and Industries. Advanced filtering and searching allow them to easily locate specific adhesives, sealants, or hardware items by SKU or Name.', style='List Bullet')
document.add_paragraph('Product Details: Each product page includes pricing, GST calculations, detailed descriptions, and real-time stock alerts. If stock falls below 50 units, an "Only X units left" warning is displayed.', style='List Bullet')
document.add_paragraph('Social Sharing: Customers can use the integrated share icons on product pages to share specific items with their network on WhatsApp, Facebook, LinkedIn, etc.', style='List Bullet')

document.add_heading('2.3 Checkout & Ordering', level=2)
document.add_paragraph('Cart: Customers can add items to their cart and view live subtotal calculations including estimated GST.', style='List Bullet')
document.add_paragraph('Checkout: Users submit their billing and shipping details to finalize their order securely.', style='List Bullet')

document.add_heading('2.4 Request For Quotation (RFQ)', level=2)
document.add_paragraph('For B2B bulk purchases, customers can use the "Request RFQ" feature. They can submit their company details, GST number, and specific quantity requirements along with a specification sheet attachment for custom pricing consideration.', style='List Bullet')

document.add_heading('2.5 Support & Tracking', level=2)
document.add_paragraph('Customers can view their Order History and live tracking statuses in their profile. Integrated floating widgets allow 1-click support via WhatsApp, Instagram, and LinkedIn.', style='List Bullet')

document.add_heading('3. Administrator Workflow (Admin Console)', level=1)

document.add_heading('3.1 Dashboard & Metrics', level=2)
document.add_paragraph('The admin can log into the secure Admin Console to view high-level metrics including total revenue, active orders, total customers, and pending RFQs.', style='List Bullet')

document.add_heading('3.2 Catalog & Inventory Management', level=2)
document.add_paragraph('Categories & Brands: Admin can create, edit, or remove product categories and brands.', style='List Bullet')
document.add_paragraph('Products: Admin has full control to add new products or edit existing ones. Key editable fields include Name, Description, Basic Price, Discount %, HSN Code, Stock Quantity, and Product Images.', style='List Bullet')
document.add_paragraph('Inventory Updates: Stock levels can be updated directly from the central dashboard. When the stock reaches zero, the item automatically updates its status on the front-end.', style='List Bullet')

document.add_heading('3.3 Order Management', level=2)
document.add_paragraph('The admin can view all placed orders, filtering by recent or specific IDs.', style='List Bullet')
document.add_paragraph('Order states can be updated systematically: Pending -> Processing -> Shipped -> Delivered.', style='List Bullet')
document.add_paragraph('When marking an order as shipped, the admin can input Courier details and a Tracking URL, which instantly becomes visible to the customer.', style='List Bullet')

document.add_heading('3.4 Promotional Management (Offers)', level=2)
document.add_paragraph('Top Banner Carousel: Admin can upload large promotional images that automatically slide at the very top of the homepage.', style='List Bullet')
document.add_paragraph('Offer Posters (Popups): Admin can configure popup images that greet the user when they first load the homepage.', style='List Bullet')
document.add_paragraph('Flash Deals & Featured Products: Admin can easily toggle the "Flash Sale" or "Featured" switch on any product to showcase it prominently on the homepage grids.', style='List Bullet')

document.add_heading('3.5 RFQ (Request For Quotation) Processing', level=2)
document.add_paragraph('The admin console includes an RFQ section to review all incoming B2B requests. The admin can download attached specification sheets, view the requested bulk quantities, and respond to the customer via the provided contact email.', style='List Bullet')

document.add_heading('3.6 Site Settings & Operations', level=2)
document.add_paragraph('Administrators can configure global variables such as standard Shipping Charges, Minimum Free Shipping Thresholds, and unified GST percentages. Email routing is integrated to notify "Ckpgrouponline@gmail.com" on relevant system events.', style='List Bullet')

document.add_page_break()

document.add_heading('4. Setup & Deployment Recommendations', level=1)
document.add_paragraph('This platform is built using React (Vite) and Supabase (PostgreSQL + Auth + Storage). To manage the live application:', style='List Bullet')
document.add_paragraph('Access Supabase: Ensure the client has the credentials to the Supabase dashboard to monitor raw database tables and manage authentication policies if needed.', style='List Bullet')
document.add_paragraph('Manage Storage: Product images, Top Banners, and Offer Posters are stored in the "aditya-assets" storage bucket. Ensure the bucket does not exceed the plan limits.', style='List Bullet')

p = document.add_paragraph('\nThank you for choosing ')
p.add_run('Zeony Technologies').bold = True
p.add_run('. If you have any further questions or require technical support, please contact our implementation team.')

document.save('Aditya_Enterprises_Project_Handover.docx')
